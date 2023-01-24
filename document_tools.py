import csv
from docx import Document
import analysis_tools as at

class DataFile:
   def __init__(self,filename,title):
       self.filename=filename
       self.title=title
       self.description= None
       self.process= []
       self.data= []
       self.header= None

   def get_filename(self):
       return self.filename

   def set_filename(self,new_name):
       self.filename = new_name

   def get_title(self) :
       return self.title

   def set_title(self,new_title):
       self.title = new_title

   def get_description (self):
       return self.description

   def set_description(self,new_desc):
       self.description = new_desc

   def get_process (self):
       return self.process

   def set_process(self,new_process):
       self.process = new_process

   def get_data(self):
       return self.data

   def set_data(self,new_data):
       self.data = new_data

   def get_header(self):
       return self.header

   def set_header(self,new_header):
       self.header=new_header


   def get_col_name(self,index):
       return self.header[index]

   def get_col(self,col):
       new_list=[]
       for row in self.data:
           new_list.append(row[col])
       return new_list

   def load(self):
       data = []
       with open('./data/vgsales.csv') as csvfile:
         reader = csv.reader(csvfile)
         for row in reader:
           data.append(row)
         self.header = data[0]
         del(data[0])
         self.data = data
class ReportGenerator:
   def __init__(self,datafile):
       self.datafile=datafile

   def generate(self, output_name):
       d = Document()
       d.add_heading(self.datafile.get_title(),0)
       d.add_paragraph(self.datafile.get_description())
       for r in self.datafile.get_process():
           d.add_heading(self.datafile.get_col_name(r),1)
           col = self.datafile.get_col(r)
           table = d.add_table(rows=1,cols=2)

           row_cells = table.add_row().cells
           row_cells[0].text = 'Mean:'
           print(at.MyMath.cleanup(self.datafile.get_col(r)))
           row_cells[1].text = str(at.MyMath.mean(self.datafile.get_col(r)))

           row_cells = table.add_row().cells
           row_cells[0].text = 'Standard Deviation:'
           row_cells[1].text = str(at.MyMath.stddev(self.datafile.get_col(r)))

           row_cells = table.add_row().cells
           row_cells[0].text = 'Median:'
           row_cells[1].text = str(at.MyMath.median(self.datafile.get_col(r)))

           row_cells = table.add_row().cells
           row_cells[0].text = 'Minimum:'
           row_cells[1].text = str(at.MyMath.min(self.datafile.get_col(r)))

           row_cells = table.add_row().cells
           row_cells[0].text = 'Maximum:'
           row_cells[1].text = str(at.MyMath.max(self.datafile.get_col(r)))
       d.save(output_name)
